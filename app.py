import io
import json
from pathlib import Path

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt

BASE = Path(__file__).parent
DATA = BASE / "data"

st.set_page_config(
    page_title="Plastic Waste Audit",
    page_icon="♻️",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
<style>
.block-container {padding-top: 1.2rem; padding-bottom: 3rem; max-width: 1380px;}
[data-testid="stMetric"] {border:1px solid rgba(128,128,128,.18); padding:14px 16px; border-radius:16px; background:rgba(128,128,128,.035);}
.hero {padding:24px 26px; border-radius:20px; background:rgba(34,139,94,.08); border:1px solid rgba(34,139,94,.20); margin-bottom:1rem;}
.hero h1 {margin:0 0 .35rem 0; font-size:2.15rem;}
.hero p {margin:0; font-size:1.02rem; opacity:.82;}
.insight {padding:16px 18px; border-left:4px solid #2e8b57; background:rgba(34,139,94,.07); border-radius:0 12px 12px 0; margin:.4rem 0 1.2rem 0;}
.explain {font-size:.94rem; opacity:.80;}
.section-gap {margin-top:1.2rem;}
</style>
""",
    unsafe_allow_html=True,
)


@st.cache_data
def load_data():
    df = pd.read_csv(DATA / "brand_audit_clean.csv")
    cleanup = pd.read_csv(DATA / "monthly_plastic_weight.csv", parse_dates=["Date"])
    with open(DATA / "data_quality_summary.json", encoding="utf-8") as f:
        quality = json.load(f)
    return df, cleanup, quality


def clean_text(series):
    return series.fillna("").astype(str).str.strip()


def identifiable_mask(frame):
    parent = clean_text(frame["Parent Company"]).str.lower()
    brand = clean_text(frame["Brand Name"]).str.lower()
    return (
        ~parent.isin(["unbranded", "unidentified / missing", "unknown", "unkown", ""])
        & ~brand.str.match(r"^(unknown|unkown|unbranded)$")
        & brand.ne("")
    )


def friendly_material(value):
    mapping = {
        "PET": "PET plastic",
        "HDPE": "HDPE plastic",
        "LDPE": "LDPE plastic",
        "PP": "Polypropylene (PP)",
        "PS": "Polystyrene (PS)",
        "PVC": "PVC plastic",
        "O": "Other material",
    }
    return mapping.get(str(value), str(value))


def friendly_layer(value):
    return {"ML": "Multi-layer packaging", "SL": "Single-layer packaging"}.get(str(value), str(value))


def moneyless_number(value):
    return f"{value:,.0f}"


def policy_brief_text(summary):
    return f"""# Policy Brief: What the Plastic Waste Audit Shows

## Purpose
This brief translates the plastic waste audit into clear findings and practical actions for decision-makers, producers, community organisations and waste-management partners.

## What was observed
The audit recorded **{summary['total_items']:,.0f} items** across **{summary['brands']} brands** and **{summary['companies']} identifiable parent companies**. Recorded cleanup activity collected **{summary['cleanup_kg']:,.1f} kg** of plastic across **{summary['cleanup_events']} cleanup events**.

A major finding is that **{summary['unidentified_share']:.1f}% of all recorded items could not be confidently linked to an identifiable producer**. Company comparisons therefore use only identifiable branded items and should be read as an audit snapshot, not as market share or total environmental responsibility.

## Key findings
- **Packaging profile:** {summary['top_layer_name']} was the most common packaging layer, accounting for **{summary['top_layer_share']:.1f}%** of recorded items.
- **Material profile:** {summary['top_material_name']} was the most common material category, accounting for **{summary['top_material_share']:.1f}%** of recorded items.
- **Concentration among identifiable producers:** the top five parent companies accounted for **{summary['top5_share']:.1f}%** of identifiable branded items; the top ten accounted for **{summary['top10_share']:.1f}%**.
- **Leading identifiable company in this audit:** {summary['top_company']} accounted for **{summary['top_company_share']:.1f}%** of identifiable branded items.
- **Cleanup record:** {summary['cleanup_kg']:,.1f} kg was collected across the recorded cleanup events. Because location and date both changed, these observations should not be interpreted as a time trend.

## What the findings mean
The audit suggests that action can be made more targeted. A relatively small set of identifiable companies accounts for a large share of the branded items observed, while multi-layer and common plastic packaging types form an important part of the waste stream. At the same time, the high share of unbranded or unidentifiable items limits producer attribution and shows why stronger audit protocols and packaging identification matter.

## Recommended actions
1. **Prioritise engagement with the most frequently observed producers.** Use the ranking as a starting point for dialogue on collection, recovery, redesign and producer-supported waste programmes.
2. **Focus recovery and prevention efforts on the dominant packaging types.** Use the material and layer profile to target the packaging formats most frequently found in the audit.
3. **Improve identification of unbranded waste.** Add clearer field protocols, photo documentation and standard coding so more items can be attributed accurately in future audits.
4. **Repeat the audit consistently over time.** Use the same sites, time windows and collection methods to make future comparisons more defensible.
5. **Link audit evidence to local waste action.** Use site-level findings to guide cleanup planning, collection partnerships, community education and producer engagement.
6. **Avoid overclaiming from a single audit.** Treat these results as observed evidence from the sampled audit, not as estimates of national market share, total company production or total environmental impact.

## Evidence limitations
The analysis is based on items recorded in the source audit. Some records are unbranded, damaged or cannot be linked confidently to a parent company. The current cleanup table contains only {summary['cleanup_events']} events, with location and date changing together. Future rounds should strengthen consistency in sampling, site coverage, coding and repeated measurement.

## Bottom line
The audit is most useful as a **targeting and accountability tool**: it shows what types of packaging are being found, which identifiable companies appear most often, and where better producer identification and repeated measurement are needed. The strongest next step is to use the findings to focus engagement and then repeat the audit using a standardised methodology.
"""


def build_policy_docx(summary):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_heading("Policy Brief: What the Plastic Waste Audit Shows", level=0)
    title.style = doc.styles["Title"]
    doc.add_paragraph("Evidence summary and practical actions for decision-makers, producers and waste-management partners.")

    doc.add_heading("At a glance", level=1)
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    labels = ["Items recorded", "Brands identified", "Companies identified", "Plastic collected"]
    values = [f"{summary['total_items']:,.0f}", str(summary['brands']), str(summary['companies']), f"{summary['cleanup_kg']:,.1f} kg"]
    for i, label in enumerate(labels):
        table.cell(0, i).text = label
        table.cell(1, i).text = values[i]

    doc.add_heading("What was observed", level=1)
    doc.add_paragraph(
        f"The audit recorded {summary['total_items']:,.0f} items across {summary['brands']} brands and "
        f"{summary['companies']} identifiable parent companies. Recorded cleanup activity collected "
        f"{summary['cleanup_kg']:,.1f} kg of plastic across {summary['cleanup_events']} cleanup events."
    )
    doc.add_paragraph(
        f"{summary['unidentified_share']:.1f}% of all recorded items could not be confidently linked to an identifiable producer. "
        "Company comparisons therefore use only identifiable branded items and should be interpreted as an audit snapshot, not market share or total environmental responsibility."
    )

    doc.add_heading("Key findings", level=1)
    findings = [
        f"{summary['top_layer_name']} was the most common packaging layer ({summary['top_layer_share']:.1f}% of recorded items).",
        f"{summary['top_material_name']} was the most common material category ({summary['top_material_share']:.1f}% of recorded items).",
        f"The top five parent companies accounted for {summary['top5_share']:.1f}% of identifiable branded items; the top ten accounted for {summary['top10_share']:.1f}%.",
        f"{summary['top_company']} was the leading identifiable company in this audit, representing {summary['top_company_share']:.1f}% of identifiable branded items.",
        f"Recorded cleanup activity collected {summary['cleanup_kg']:,.1f} kg across {summary['cleanup_events']} events."
    ]
    for item in findings:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Policy implications", level=1)
    doc.add_paragraph(
        "The results can support more targeted action. A relatively small set of identifiable companies accounts for a large share of the branded items observed, while the packaging profile shows where prevention, redesign, collection and recovery efforts could be focused. The high share of unbranded or unidentifiable items also limits producer attribution and should be addressed in future audit rounds."
    )

    doc.add_heading("Recommended actions", level=1)
    recommendations = [
        "Prioritise engagement with the most frequently observed producers on collection, recovery, packaging redesign and producer-supported waste programmes.",
        "Focus prevention and recovery efforts on the packaging formats and material categories most frequently found in the audit.",
        "Improve identification of unbranded waste through clearer field protocols, photo documentation and standard coding.",
        "Repeat the audit consistently over time using comparable sites, timing and collection methods.",
        "Link site-level evidence to cleanup planning, collection partnerships, community education and producer engagement.",
        "Treat the findings as observed audit evidence rather than estimates of market share, total production or total environmental impact."
    ]
    for rec in recommendations:
        doc.add_paragraph(rec, style="List Number")

    doc.add_heading("Evidence limitations", level=1)
    doc.add_paragraph(
        f"Some recorded items are unbranded, damaged or cannot be confidently linked to a parent company. The cleanup table contains only {summary['cleanup_events']} events, with location and date changing together. Future rounds should strengthen sampling consistency, site coverage, coding and repeated measurement."
    )

    doc.add_heading("Bottom line", level=1)
    doc.add_paragraph(
        "Use this audit as a targeting and accountability tool: identify what packaging is being found, which identifiable companies appear most often, and where better producer identification and repeated measurement are needed."
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


df, cleanup, quality = load_data()
df["Total Count"] = pd.to_numeric(df["Total Count"], errors="coerce").fillna(0)
df["Is Valid?"] = clean_text(df["Is Valid?"])
valid = df[(df["Is Valid?"].str.lower() == "yes") & (df["Total Count"] > 0)].copy()

# Friendly labels used only for display; original source codes stay in the data.
valid["Plastic type"] = valid["Type Material"].map(friendly_material)
valid["Packaging type"] = valid["Layers"].map(friendly_layer)

st.markdown(
    '<div class="hero"><h1>♻️ Plastic Waste Audit</h1><p>A simple view of what was found, which brands and companies appeared most often, what kinds of packaging were recorded, and what the evidence can support.</p></div>',
    unsafe_allow_html=True,
)

# Simple public-facing filters.
with st.sidebar:
    st.header("Explore the audit")
    st.caption("Choose a filter only if you want to focus on one part of the data.")

    company_options = ["All companies"] + sorted(clean_text(valid["Parent Company"])[clean_text(valid["Parent Company"]) != ""].unique().tolist())
    company_choice = st.selectbox("Company", company_options)

    brand_base = valid if company_choice == "All companies" else valid[clean_text(valid["Parent Company"]) == company_choice]
    brand_options = ["All brands"] + sorted(clean_text(brand_base["Brand Name"])[clean_text(brand_base["Brand Name"]) != ""].unique().tolist())
    brand_choice = st.selectbox("Brand", brand_options)

    material_options = ["All plastic / material types"] + sorted(valid["Plastic type"].dropna().unique().tolist())
    material_choice = st.selectbox("Plastic / material type", material_options)

    filt = valid.copy()
    if company_choice != "All companies":
        filt = filt[clean_text(filt["Parent Company"]) == company_choice]
    if brand_choice != "All brands":
        filt = filt[clean_text(filt["Brand Name"]) == brand_choice]
    if material_choice != "All plastic / material types":
        filt = filt[filt["Plastic type"] == material_choice]

    st.divider()
    st.caption("Need the underlying rows?")
    st.download_button(
        "Download current data",
        data=filt.to_csv(index=False).encode("utf-8"),
        file_name="plastic_audit_filtered.csv",
        mime="text/csv",
        use_container_width=True,
    )

all_items = valid["Total Count"].sum()
identified = valid[identifiable_mask(valid)].copy()
identified_total = identified["Total Count"].sum()
unidentified_total = all_items - identified_total
unidentified_share = unidentified_total / all_items * 100 if all_items else 0

comp_all = identified.groupby("Parent Company", as_index=False)["Total Count"].sum().sort_values("Total Count", ascending=False)
if not comp_all.empty:
    comp_all["Share %"] = comp_all["Total Count"] / comp_all["Total Count"].sum() * 100
    comp_all["Cumulative %"] = comp_all["Share %"].cumsum()
    top5_share = comp_all.head(5)["Share %"].sum()
    top10_share = comp_all.head(10)["Share %"].sum()
    top_company = comp_all.iloc[0]["Parent Company"]
    top_company_share = comp_all.iloc[0]["Share %"]
else:
    top5_share = top10_share = top_company_share = 0
    top_company = "No identifiable company"

layer_totals = valid.groupby("Packaging type")["Total Count"].sum().sort_values(ascending=False)
material_totals = valid.groupby("Plastic type")["Total Count"].sum().sort_values(ascending=False)
top_layer_name = layer_totals.index[0] if len(layer_totals) else "Not available"
top_layer_share = layer_totals.iloc[0] / all_items * 100 if len(layer_totals) and all_items else 0
top_material_name = material_totals.index[0] if len(material_totals) else "Not available"
top_material_share = material_totals.iloc[0] / all_items * 100 if len(material_totals) and all_items else 0

summary = {
    "total_items": all_items,
    "brands": valid["Brand Name"].nunique(),
    "companies": identified["Parent Company"].nunique(),
    "cleanup_kg": cleanup["Plastic Weight (Kgs)"].sum(),
    "cleanup_events": len(cleanup),
    "unidentified_share": unidentified_share,
    "top_layer_name": top_layer_name,
    "top_layer_share": top_layer_share,
    "top_material_name": top_material_name,
    "top_material_share": top_material_share,
    "top5_share": top5_share,
    "top10_share": top10_share,
    "top_company": top_company,
    "top_company_share": top_company_share,
}

pages = st.tabs([
    "Overview",
    "Companies & brands",
    "Plastic types",
    "Cleanup activity",
    "Policy brief",
    "About the audit",
])

with pages[0]:
    st.subheader("What did the audit find?")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Items recorded", f"{filt['Total Count'].sum():,.0f}")
    c2.metric("Brands represented", f"{filt['Brand Name'].nunique():,}")
    c3.metric("Companies represented", f"{filt.loc[identifiable_mask(filt), 'Parent Company'].nunique():,}")
    c4.metric("Plastic collected", f"{cleanup['Plastic Weight (Kgs)'].sum():,.1f} kg")

    st.markdown(
        f'<div class="insight"><b>Key message:</b> In the full audit, {unidentified_share:.1f}% of recorded items could not be confidently linked to an identifiable producer. Company rankings therefore focus only on identifiable branded items.</div>',
        unsafe_allow_html=True,
    )

    left, right = st.columns(2)
    ranked = filt[identifiable_mask(filt)].copy()

    top_brands = (
        ranked.groupby("Brand Name", as_index=False)["Total Count"]
        .sum().nlargest(10, "Total Count").sort_values("Total Count")
    )
    if len(top_brands):
        fig = px.bar(top_brands, x="Total Count", y="Brand Name", orientation="h", text_auto=".3s")
        fig.update_layout(title="Which brands appeared most often?", xaxis_title="Number of items found", yaxis_title=None, height=430)
        left.plotly_chart(fig, use_container_width=True)
    else:
        left.info("No identifiable brands match the current filters.")

    top_companies = (
        ranked.groupby("Parent Company", as_index=False)["Total Count"]
        .sum().nlargest(10, "Total Count").sort_values("Total Count")
    )
    if len(top_companies):
        fig = px.bar(top_companies, x="Total Count", y="Parent Company", orientation="h", text_auto=".3s")
        fig.update_layout(title="Which companies appeared most often?", xaxis_title="Number of identifiable items found", yaxis_title=None, height=430)
        right.plotly_chart(fig, use_container_width=True)
    else:
        right.info("No identifiable companies match the current filters.")

    st.subheader("What kinds of packaging were found?")
    left, right = st.columns(2)
    mat = filt.groupby("Plastic type", as_index=False)["Total Count"].sum().sort_values("Total Count", ascending=False)
    fig = px.bar(mat, x="Plastic type", y="Total Count", text_auto=".3s")
    fig.update_layout(title="Material / plastic type", xaxis_title=None, yaxis_title="Number of items found")
    left.plotly_chart(fig, use_container_width=True)

    lay = filt.groupby("Packaging type", as_index=False)["Total Count"].sum().sort_values("Total Count", ascending=False)
    fig = px.pie(lay, names="Packaging type", values="Total Count", hole=.48)
    fig.update_layout(title="Single-layer or multi-layer packaging?")
    right.plotly_chart(fig, use_container_width=True)

    if filt["Total Count"].sum():
        local_layer = lay.iloc[0]["Packaging type"] if len(lay) else "Not available"
        local_layer_share = lay.iloc[0]["Total Count"] / filt["Total Count"].sum() * 100 if len(lay) else 0
        st.markdown(
            f'<div class="insight"><b>What this means:</b> {local_layer} is the largest packaging-layer category in the current view, representing about {local_layer_share:.1f}% of recorded items.</div>',
            unsafe_allow_html=True,
        )

with pages[1]:
    st.subheader("Who appears most often among identifiable branded items?")
    st.caption("This page excludes items that could not be confidently linked to a producer.")

    producer = filt[identifiable_mask(filt)].copy()
    comp = producer.groupby("Parent Company", as_index=False)["Total Count"].sum().sort_values("Total Count", ascending=False)
    if len(comp):
        comp["Share of identifiable items (%)"] = comp["Total Count"] / comp["Total Count"].sum() * 100
        comp["Cumulative share (%)"] = comp["Share of identifiable items (%)"].cumsum()
        top5 = comp.head(5)["Share of identifiable items (%)"].sum()
        top10 = comp.head(10)["Share of identifiable items (%)"].sum()

        st.markdown(
            f'<div class="insight"><b>Key message:</b> The top five companies account for {top5:.1f}% of identifiable branded items in the current view. The top ten account for {top10:.1f}%.</div>',
            unsafe_allow_html=True,
        )

        shown = comp.head(15).copy().sort_values("Total Count")
        fig = px.bar(shown, x="Total Count", y="Parent Company", orientation="h", text_auto=".3s")
        fig.update_layout(title="Companies associated with the most identifiable items", xaxis_title="Number of items found", yaxis_title=None, height=540)
        st.plotly_chart(fig, use_container_width=True)

        with st.expander("See the concentration curve"):
            shown2 = comp.head(min(25, len(comp))).copy()
            fig = go.Figure()
            fig.add_bar(x=shown2["Parent Company"], y=shown2["Total Count"], name="Items found", yaxis="y")
            fig.add_scatter(x=shown2["Parent Company"], y=shown2["Cumulative share (%)"], name="Running share", yaxis="y2", mode="lines+markers")
            fig.update_layout(
                title="How quickly the observed branded items concentrate among companies",
                height=560,
                xaxis_tickangle=-55,
                yaxis=dict(title="Items found"),
                yaxis2=dict(title="Running share (%)", overlaying="y", side="right", range=[0, 105]),
                legend=dict(orientation="h", y=1.08),
            )
            st.plotly_chart(fig, use_container_width=True)
            st.caption("The line shows how much of the identifiable branded waste is accounted for as companies are added from highest to lowest count.")

        st.subheader("Look at one company")
        company_choices = comp["Parent Company"].tolist()
        selected_company = st.selectbox("Choose a company", company_choices, key="company_profile")
        d = producer[producer["Parent Company"] == selected_company]
        company_share = d["Total Count"].sum() / producer["Total Count"].sum() * 100 if producer["Total Count"].sum() else 0
        c1, c2, c3 = st.columns(3)
        c1.metric("Items linked to this company", f"{d['Total Count'].sum():,.0f}")
        c2.metric("Brands found", f"{d['Brand Name'].nunique():,}")
        c3.metric("Share of identifiable items", f"{company_share:.1f}%")
        if company_share >= 10:
            plain = f"About {round(company_share/10):.0f} in every 10 identifiable branded items in this view were linked to {selected_company}."
        else:
            plain = f"{selected_company} represents {company_share:.1f}% of identifiable branded items in this view."
        st.info(plain)

        l, r = st.columns(2)
        brand_breakdown = d.groupby("Brand Name", as_index=False)["Total Count"].sum().sort_values("Total Count", ascending=False)
        l.plotly_chart(px.bar(brand_breakdown, x="Brand Name", y="Total Count", title="Brands linked to this company"), use_container_width=True)
        material_breakdown = d.groupby("Plastic type", as_index=False)["Total Count"].sum()
        r.plotly_chart(px.pie(material_breakdown, names="Plastic type", values="Total Count", hole=.45, title="What materials were found?"), use_container_width=True)
    else:
        st.info("No identifiable company records match the current filters.")

with pages[2]:
    st.subheader("What types of plastic and packaging are being found?")

    mat = filt.groupby("Plastic type", as_index=False)["Total Count"].sum().sort_values("Total Count", ascending=False)
    lay = filt.groupby("Packaging type", as_index=False)["Total Count"].sum().sort_values("Total Count", ascending=False)
    total_view = filt["Total Count"].sum()

    if total_view and len(mat):
        material_message = f"{mat.iloc[0]['Plastic type']} is the largest material category in the current view, accounting for {mat.iloc[0]['Total Count']/total_view*100:.1f}% of items."
    else:
        material_message = "No material records match the current filters."
    st.markdown(f'<div class="insight"><b>Key message:</b> {material_message}</div>', unsafe_allow_html=True)

    left, right = st.columns(2)
    fig = px.bar(mat, x="Plastic type", y="Total Count", text_auto=".3s")
    fig.update_layout(title="Items by material / plastic type", xaxis_title=None, yaxis_title="Number of items found")
    left.plotly_chart(fig, use_container_width=True)

    fig = px.bar(lay, x="Packaging type", y="Total Count", text_auto=".3s")
    fig.update_layout(title="Items by packaging layer", xaxis_title=None, yaxis_title="Number of items found")
    right.plotly_chart(fig, use_container_width=True)

    st.subheader("Which companies are associated with which materials?")
    st.caption("This view is limited to identifiable producer records.")
    material_company = filt[identifiable_mask(filt)].copy()
    cross = material_company.pivot_table(index="Parent Company", columns="Plastic type", values="Total Count", aggfunc="sum", fill_value=0)
    if not cross.empty:
        topnames = material_company.groupby("Parent Company")["Total Count"].sum().nlargest(15).index
        cross = cross.reindex(topnames).fillna(0)
        fig = px.imshow(cross, aspect="auto", labels=dict(x="Material / plastic type", y="Company", color="Items found"))
        fig.update_layout(title="Material profile of the most frequently observed companies", height=590)
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("No identifiable company records match the current filters.")

with pages[3]:
    st.subheader("How much plastic was collected?")
    c1, c2, c3 = st.columns(3)
    c1.metric("Recorded cleanup events", f"{len(cleanup):,}")
    c2.metric("Plastic collected", f"{cleanup['Plastic Weight (Kgs)'].sum():,.1f} kg")
    c3.metric("Sacks collected", f"{cleanup['No. of sacks'].sum():,.1f}")

    st.markdown(
        '<div class="insight"><b>Important:</b> The workbook contains only three cleanup events, and the location changes along with the date. The chart below shows what was recorded at each event; it is not evidence of an increasing or decreasing time trend.</div>',
        unsafe_allow_html=True,
    )

    clean_display = cleanup.copy()
    clean_display["Date label"] = clean_display["Date"].dt.strftime("%d %b %Y")
    fig = px.bar(
        clean_display.sort_values("Date"),
        x="Location",
        y="Plastic Weight (Kgs)",
        text_auto=".1f",
        hover_data=["Date label", "No. of sacks", "Kg per Sack"],
    )
    fig.update_layout(title="Plastic collected at each recorded cleanup", xaxis_title=None, yaxis_title="Plastic collected (kg)")
    st.plotly_chart(fig, use_container_width=True)

    table = clean_display[["Location", "Date label", "Plastic Weight (Kgs)", "No. of sacks", "Kg per Sack"]].copy()
    table.columns = ["Location", "Date", "Plastic collected (kg)", "Sacks", "Kg per sack"]
    table["Kg per sack"] = table["Kg per sack"].round(1)
    st.dataframe(table, use_container_width=True, hide_index=True)

with pages[4]:
    st.subheader("Policy brief")
    st.write("This brief is generated from the same cleaned dataset as the dashboard, so the headline figures remain aligned with the evidence shown here.")

    c1, c2, c3 = st.columns(3)
    c1.metric("Observed items", f"{summary['total_items']:,.0f}")
    c2.metric("Unidentified producer share", f"{summary['unidentified_share']:.1f}%")
    c3.metric("Top 5 company share*", f"{summary['top5_share']:.1f}%")
    st.caption("*Share of identifiable branded items only.")

    st.markdown(policy_brief_text(summary))

    docx_bytes = build_policy_docx(summary)
    col1, col2 = st.columns(2)
    col1.download_button(
        "Download policy brief (Word)",
        data=docx_bytes,
        file_name="plastic_waste_audit_policy_brief.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
    col2.download_button(
        "Download policy brief (Markdown)",
        data=policy_brief_text(summary).encode("utf-8"),
        file_name="plastic_waste_audit_policy_brief.md",
        mime="text/markdown",
        use_container_width=True,
    )

    st.info("For external publication, add the organisation name, audit period, geographic scope, methodology, contact details and any policy/legal context you want the brief to address.")

with pages[5]:
    st.subheader("About this audit")
    st.write(
        "This dashboard summarises records from the source plastic brand audit. It is designed to make the evidence easier to understand without changing the original item counts."
    )

    c1, c2, c3 = st.columns(3)
    c1.metric("Source rows", f"{quality['source_rows']:,}")
    c2.metric("Valid rows used by default", f"{quality['valid_rows']:,}")
    c3.metric("Rows missing parent company", f"{quality['missing_parent_rows']:,}")

    st.markdown("""
### How to read the results
- Overall totals include valid positive-count audit records.
- Company rankings use only items that can be linked to an identifiable brand and parent company.
- Items that are unbranded or cannot be confidently attributed are **not guessed** into a company.
- Friendly labels such as “PET plastic” and “Multi-layer packaging” are used for readability, while the original source codes remain in the downloadable data.
- The item counts show what was observed in this audit. They are **not automatically equivalent to market share, production volume, or total environmental responsibility**.

### Why some items are unidentified
Packaging can be damaged, unbranded, fragmented, faded or recorded without enough information to link it confidently to a producer. Keeping these items separate is more defensible than forcing an attribution.

### How to make future audits stronger
Use the same sites and collection windows, standardise item coding, photograph hard-to-identify packaging, record location consistently, and repeat the audit over time. That will make comparisons and policy conclusions more reliable.
""")

    with st.expander("See records that may need review"):
        issues = df[(df["Is Valid?"].str.lower() != "yes") | (clean_text(df["Parent Company"]).isin(["", "Unidentified / missing"])) | (df["Total Count"] <= 0)]
        st.dataframe(issues, use_container_width=True, hide_index=True)
